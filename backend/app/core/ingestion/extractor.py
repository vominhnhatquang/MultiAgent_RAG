"""Extract raw text + page metadata from uploaded files."""
import io
from dataclasses import dataclass, field

import structlog

from app.exceptions import AppError

logger = structlog.get_logger(__name__)


@dataclass
class ExtractedPage:
    page_number: int
    text: str


@dataclass
class ExtractedDocument:
    pages: list[ExtractedPage] = field(default_factory=list)
    language: str = "vi"  # default; enricher may detect


async def extract(file_bytes: bytes, filename: str) -> ExtractedDocument:
    """Dispatch to the correct extractor based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower()
    logger.info("extractor.start", filename=filename, ext=ext, size=len(file_bytes))

    if ext == "pdf":
        return _extract_pdf(file_bytes)
    elif ext == "docx":
        return _extract_docx(file_bytes)
    elif ext in ("md", "txt"):
        return _extract_text(file_bytes)
    else:
        raise AppError(f"Unsupported file type: {ext}", "UNSUPPORTED_TYPE", 415)


def _extract_pdf(data: bytes) -> ExtractedDocument:
    try:
        import fitz  # pymupdf
    except ImportError as exc:
        raise AppError("pymupdf not installed", "INTERNAL_ERROR", 500) from exc

    doc_out = ExtractedDocument()
    with fitz.open(stream=data, filetype="pdf") as pdf:
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text("text").strip()
            if text:
                doc_out.pages.append(ExtractedPage(page_number=page_num, text=text))
    logger.info("extractor.pdf_done", pages=len(doc_out.pages))
    return doc_out


def _extract_docx(data: bytes) -> ExtractedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:
        raise AppError("python-docx not installed", "INTERNAL_ERROR", 500) from exc

    doc_out = ExtractedDocument()
    docx = DocxDocument(io.BytesIO(data))
    full_text = "\n".join(p.text for p in docx.paragraphs if p.text.strip())
    doc_out.pages.append(ExtractedPage(page_number=1, text=full_text))
    logger.info("extractor.docx_done", chars=len(full_text))
    return doc_out


def _extract_text(data: bytes) -> ExtractedDocument:
    text = data.decode("utf-8", errors="replace").strip()
    doc_out = ExtractedDocument()
    doc_out.pages.append(ExtractedPage(page_number=1, text=text))
    logger.info("extractor.text_done", chars=len(text))
    return doc_out
