#!/usr/bin/env python3
"""
Standalone ingestion script — không cần Docker, PostgreSQL, hay Qdrant.

Chạy toàn bộ pipeline:
  data/uploads/raw/*.pdf  →  Extract → Clean → Chunk → Enrich → Embed (Ollama)
                          →  Lưu vào  data/processed/chunks.db  (SQLite)
                                      data/processed/vectors.jsonl (vectors)

Yêu cầu:
  - Ollama đang chạy:  ollama serve  (cửa sổ terminal khác)
  - Model đã pull:     ollama pull nomic-embed-text
  - Python packages:   pip install pymupdf tiktoken httpx --break-system-packages

Chạy:
  python3 scripts/ingest_local.py
  python3 scripts/ingest_local.py --raw-dir /path/to/files --embed-model nomic-embed-text
"""

import argparse
import hashlib
import json
import logging
import re
import sqlite3
import sys
import unicodedata
import uuid
from dataclasses import asdict, dataclass, field
from datetime import datetime, timezone
from pathlib import Path

import httpx

# ─── Config ────────────────────────────────────────────────────────────────────

CHUNK_SIZE    = 512   # tokens
CHUNK_OVERLAP = 64    # tokens
EMBED_BATCH   = 16    # chunks per Ollama request

OLLAMA_BASE   = "http://localhost:11434"
EMBED_MODEL   = "nomic-embed-text"
TIMEOUT       = 120   # seconds

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("ingest")


# ─── Data classes ──────────────────────────────────────────────────────────────

@dataclass
class ExtractedPage:
    page_number: int
    text: str


@dataclass
class ExtractedDocument:
    pages: list[ExtractedPage] = field(default_factory=list)


@dataclass
class TextChunk:
    content: str
    chunk_index: int
    page_number: int | None
    token_count: int
    char_count: int
    metadata: dict = field(default_factory=dict)


# ─── Step 1: Extract ───────────────────────────────────────────────────────────

def extract(file_bytes: bytes, filename: str) -> ExtractedDocument:
    ext = filename.rsplit(".", 1)[-1].lower()
    log.info(f"[extract] {filename}  ({len(file_bytes):,} bytes)")

    if ext == "pdf":
        return _extract_pdf(file_bytes)
    elif ext == "docx":
        return _extract_docx(file_bytes)
    elif ext in ("md", "txt"):
        return _extract_text(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")


def _extract_pdf(data: bytes) -> ExtractedDocument:
    import fitz  # pymupdf
    doc_out = ExtractedDocument()
    with fitz.open(stream=data, filetype="pdf") as pdf:
        for page_num, page in enumerate(pdf, start=1):
            text = page.get_text("text").strip()
            if text:
                doc_out.pages.append(ExtractedPage(page_number=page_num, text=text))
    log.info(f"[extract] PDF → {len(doc_out.pages)} trang có text")
    return doc_out


def _extract_docx(data: bytes) -> ExtractedDocument:
    import io
    from docx import Document as DocxDocument
    doc_out = ExtractedDocument()
    docx = DocxDocument(io.BytesIO(data))
    full_text = "\n".join(p.text for p in docx.paragraphs if p.text.strip())
    doc_out.pages.append(ExtractedPage(page_number=1, text=full_text))
    return doc_out


def _extract_text(data: bytes) -> ExtractedDocument:
    text = data.decode("utf-8", errors="replace").strip()
    doc_out = ExtractedDocument()
    doc_out.pages.append(ExtractedPage(page_number=1, text=text))
    return doc_out


# ─── Step 2: Clean ─────────────────────────────────────────────────────────────

def clean(doc: ExtractedDocument) -> ExtractedDocument:
    cleaned_pages = []
    for page in doc.pages:
        text = page.text
        # Unicode NFKC normalize
        text = unicodedata.normalize("NFC", text)
        # Xoá control characters (giữ \n \t)
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        # Collapse whitespace trong từng dòng
        text = "\n".join(
            re.sub(r"[ \t]+", " ", line).strip()
            for line in text.split("\n")
        )
        # Xoá dòng chỉ chứa số (page numbers)
        text = re.sub(r"(?m)^\d{1,4}$", "", text)
        # Collapse ≥3 xuống dòng liên tiếp
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
        if text:
            cleaned_pages.append(ExtractedPage(page_number=page.page_number, text=text))

    log.info(f"[clean] {len(doc.pages)} trang → {len(cleaned_pages)} trang sau clean")
    return ExtractedDocument(pages=cleaned_pages)


# ─── Step 3: Chunk ─────────────────────────────────────────────────────────────

def _count_tokens(text: str) -> int:
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(text))
    except Exception:
        return max(1, len(text) // 4)


def _split_into_sentences(text: str) -> list[str]:
    parts = re.split(r"(?<=[.!?।\n])\s+", text)
    return [p.strip() for p in parts if p.strip()]


def _make_chunk(content: str, index: int, page: int | None) -> TextChunk:
    return TextChunk(
        content=content,
        chunk_index=index,
        page_number=page,
        token_count=_count_tokens(content),
        char_count=len(content),
    )


def _overlap_sentences(sentences: list[str], overlap_tokens: int) -> list[str]:
    result: list[str] = []
    total = 0
    for s in reversed(sentences):
        t = _count_tokens(s)
        if total + t > overlap_tokens:
            break
        result.insert(0, s)
        total += t
    return result


def _overlap_words(words: list[str], overlap_tokens: int) -> list[str]:
    result: list[str] = []
    total = 0
    for w in reversed(words):
        t = _count_tokens(w)
        if total + t > overlap_tokens:
            break
        result.insert(0, w)
        total += t
    return result


def chunk_document(doc: ExtractedDocument) -> list[TextChunk]:
    chunks: list[TextChunk] = []
    chunk_index = 0

    for page in doc.pages:
        sentences = _split_into_sentences(page.text)
        buffer: list[str] = []
        buffer_tokens = 0

        for sentence in sentences:
            s_tokens = _count_tokens(sentence)

            if s_tokens > CHUNK_SIZE:
                words = sentence.split()
                sub_buf: list[str] = []
                sub_tok = 0
                for word in words:
                    w_tok = _count_tokens(word)
                    if sub_tok + w_tok > CHUNK_SIZE and sub_buf:
                        content = " ".join(sub_buf)
                        chunks.append(_make_chunk(content, chunk_index, page.page_number))
                        chunk_index += 1
                        keep = _overlap_words(sub_buf, CHUNK_OVERLAP)
                        sub_buf = keep
                        sub_tok = _count_tokens(" ".join(sub_buf))
                    sub_buf.append(word)
                    sub_tok += w_tok
                if sub_buf:
                    buffer = sub_buf
                    buffer_tokens = sub_tok
                continue

            if buffer_tokens + s_tokens > CHUNK_SIZE and buffer:
                content = " ".join(buffer)
                chunks.append(_make_chunk(content, chunk_index, page.page_number))
                chunk_index += 1
                buffer = _overlap_sentences(buffer, CHUNK_OVERLAP)
                buffer_tokens = _count_tokens(" ".join(buffer))

            buffer.append(sentence)
            buffer_tokens += s_tokens

        if buffer:
            content = " ".join(buffer)
            chunks.append(_make_chunk(content, chunk_index, page.page_number))
            chunk_index += 1

    log.info(f"[chunk] {len(chunks)} chunks")
    return chunks


# ─── Step 4: Enrich ────────────────────────────────────────────────────────────

_VI_CHARS = set("àáảãạăắặẳẵặâấầẩẫậèéẻẽẹêếềểễệìíỉĩịòóỏõọôốồổỗộơớờởỡợùúủũụưứừửữựỳýỷỹỵđ"
               "ÀÁẢÃẠĂẮẶẲẴẶÂẤẦẨẪẬÈÉẺẼẸÊẾỀỂỄỆÌÍỈĨỊÒÓỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÙÚỦŨỤƯỨỪỬỮỰỲÝỶỸỴĐ")

_STOPWORDS = {
    "và", "của", "là", "có", "trong", "được", "cho", "với", "các", "một",
    "the", "is", "in", "of", "and", "a", "an", "to", "for", "that", "this",
    "it", "as", "at", "by", "be", "are", "was", "were", "from", "or", "on",
}


def _detect_language(text: str) -> str:
    vi_count = sum(1 for c in text if c in _VI_CHARS)
    return "vi" if vi_count > 5 else "en"


def _extract_keywords(text: str, top_n: int = 5) -> list[str]:
    words = re.findall(r"\b[a-zA-ZÀ-ỹ]{3,}\b", text.lower())
    freq: dict[str, int] = {}
    for w in words:
        if w not in _STOPWORDS:
            freq[w] = freq.get(w, 0) + 1
    top = sorted(freq, key=lambda x: -freq[x])[:top_n]
    return top


def enrich(chunks: list[TextChunk]) -> list[TextChunk]:
    for ch in chunks:
        ch.metadata["language"] = _detect_language(ch.content)
        ch.metadata["keywords"] = _extract_keywords(ch.content)
    log.info(f"[enrich] {len(chunks)} chunks enriched")
    return chunks


# ─── Step 5: Embed ─────────────────────────────────────────────────────────────

def embed(chunks: list[TextChunk], model: str, base_url: str) -> list[tuple[TextChunk, list[float]]]:
    results: list[tuple[TextChunk, list[float]]] = []
    total = len(chunks)

    with httpx.Client(timeout=TIMEOUT) as client:
        for i in range(0, total, EMBED_BATCH):
            batch = chunks[i : i + EMBED_BATCH]
            texts = [c.content for c in batch]

            log.info(f"[embed] batch {i//EMBED_BATCH + 1}/{(total-1)//EMBED_BATCH + 1}"
                     f"  ({len(batch)} chunks)")

            resp = client.post(
                f"{base_url}/api/embed",
                json={"model": model, "input": texts},
            )
            resp.raise_for_status()
            vectors = resp.json()["embeddings"]

            for chunk_obj, vec in zip(batch, vectors):
                results.append((chunk_obj, vec))

    log.info(f"[embed] done — {len(results)} vectors ({len(results[0][1])}d)")
    return results


# ─── Step 6: Index → SQLite + JSONL ────────────────────────────────────────────

def init_db(db_path: Path) -> sqlite3.Connection:
    conn = sqlite3.connect(db_path)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS documents (
            id          TEXT PRIMARY KEY,
            filename    TEXT NOT NULL,
            file_hash   TEXT UNIQUE NOT NULL,
            file_type   TEXT NOT NULL,
            file_size   INTEGER NOT NULL,
            status      TEXT NOT NULL DEFAULT 'processing',
            chunk_count INTEGER DEFAULT 0,
            created_at  TEXT NOT NULL,
            updated_at  TEXT NOT NULL
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS chunks (
            id          TEXT PRIMARY KEY,
            document_id TEXT NOT NULL REFERENCES documents(id),
            chunk_index INTEGER NOT NULL,
            page_number INTEGER,
            content     TEXT NOT NULL,
            token_count INTEGER,
            char_count  INTEGER,
            metadata    TEXT,
            created_at  TEXT NOT NULL
        )
    """)
    conn.commit()
    return conn


def index_to_db(
    doc_id: str,
    filename: str,
    chunk_vectors: list[tuple[TextChunk, list[float]]],
    conn: sqlite3.Connection,
    vectors_path: Path,
) -> int:
    now = datetime.now(timezone.utc).isoformat()

    # Append vectors to JSONL
    with vectors_path.open("a", encoding="utf-8") as vf:
        for ch, vec in chunk_vectors:
            chunk_id = str(uuid.uuid4())
            # Write chunk record
            conn.execute("""
                INSERT OR REPLACE INTO chunks
                  (id, document_id, chunk_index, page_number, content,
                   token_count, char_count, metadata, created_at)
                VALUES (?,?,?,?,?,?,?,?,?)
            """, (
                chunk_id, doc_id, ch.chunk_index, ch.page_number,
                ch.content, ch.token_count, ch.char_count,
                json.dumps(ch.metadata, ensure_ascii=False), now,
            ))
            # Write vector
            vf.write(json.dumps({
                "id": chunk_id,
                "doc_id": doc_id,
                "filename": filename,
                "chunk_index": ch.chunk_index,
                "page_number": ch.page_number,
                "language": ch.metadata.get("language"),
                "content": ch.content,
                "vector": vec,
            }, ensure_ascii=False) + "\n")

    conn.commit()
    return len(chunk_vectors)


# ─── Main pipeline ─────────────────────────────────────────────────────────────

def ingest_file(
    file_path: Path,
    conn: sqlite3.Connection,
    vectors_path: Path,
    model: str,
    base_url: str,
) -> dict:
    filename = file_path.name
    file_bytes = file_path.read_bytes()
    file_hash = hashlib.sha256(file_bytes).hexdigest()
    now = datetime.now(timezone.utc).isoformat()
    doc_id = str(uuid.uuid4())

    # Check duplicate
    row = conn.execute(
        "SELECT id, status FROM documents WHERE file_hash = ?", (file_hash,)
    ).fetchone()
    if row:
        log.warning(f"[skip] {filename} đã tồn tại (doc_id={row[0]}, status={row[1]})")
        return {"file": filename, "status": "duplicate", "doc_id": row[0]}

    # Register
    conn.execute("""
        INSERT INTO documents (id, filename, file_hash, file_type, file_size, status, created_at, updated_at)
        VALUES (?,?,?,?,?,?,?,?)
    """, (doc_id, filename, file_hash, filename.rsplit(".",1)[-1].lower(),
          len(file_bytes), "processing", now, now))
    conn.commit()

    try:
        extracted  = extract(file_bytes, filename)
        cleaned    = clean(extracted)

        if not cleaned.pages:
            raise ValueError("Document rỗng sau khi clean")

        chunks     = chunk_document(cleaned)
        enriched   = enrich(chunks)
        embeddings = embed(enriched, model, base_url)
        count      = index_to_db(doc_id, filename, embeddings, conn, vectors_path)

        conn.execute(
            "UPDATE documents SET status='indexed', chunk_count=?, updated_at=? WHERE id=?",
            (count, datetime.now(timezone.utc).isoformat(), doc_id)
        )
        conn.commit()
        log.info(f"[done] {filename} → {count} chunks  doc_id={doc_id}")
        return {"file": filename, "status": "indexed", "doc_id": doc_id, "chunks": count}

    except Exception as exc:
        log.error(f"[error] {filename}: {exc}")
        conn.execute(
            "UPDATE documents SET status='error', updated_at=? WHERE id=?",
            (datetime.now(timezone.utc).isoformat(), doc_id)
        )
        conn.commit()
        return {"file": filename, "status": "error", "error": str(exc)}


# ─── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Standalone RAG ingestion (no Docker needed)")
    parser.add_argument("--raw-dir",    default="data/uploads/raw",      help="Thư mục chứa file input")
    parser.add_argument("--out-dir",    default="data/processed",        help="Thư mục output")
    parser.add_argument("--embed-model",default=EMBED_MODEL,             help="Ollama model tên")
    parser.add_argument("--ollama-url", default=OLLAMA_BASE,             help="Ollama base URL")
    parser.add_argument("--file",       default=None,                    help="Chỉ xử lý 1 file cụ thể")
    args = parser.parse_args()

    raw_dir  = Path(args.raw_dir)
    out_dir  = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    db_path      = out_dir / "chunks.db"
    vectors_path = out_dir / "vectors.jsonl"

    # Check Ollama
    log.info(f"Kiểm tra Ollama tại {args.ollama_url} ...")
    try:
        resp = httpx.get(f"{args.ollama_url}/api/tags", timeout=5)
        models = [m["name"] for m in resp.json().get("models", [])]
        log.info(f"Ollama models: {models}")
        if not any(args.embed_model in m for m in models):
            log.error(f"Model '{args.embed_model}' chưa được pull!")
            log.error(f"  Chạy:  ollama pull {args.embed_model}")
            sys.exit(1)
    except Exception as e:
        log.error(f"Không kết nối được Ollama: {e}")
        log.error("  Chạy:  ollama serve   (trong cửa sổ terminal khác)")
        sys.exit(1)

    conn = init_db(db_path)

    # Collect files
    if args.file:
        files = [Path(args.file)]
    else:
        files = sorted([
            f for f in raw_dir.iterdir()
            if f.suffix.lower() in (".pdf", ".docx", ".md", ".txt")
            and not f.name.endswith(":Zone.Identifier")
        ])

    if not files:
        log.warning(f"Không có file nào trong {raw_dir}")
        sys.exit(0)

    log.info(f"Tìm thấy {len(files)} file(s) cần xử lý")
    log.info(f"DB:      {db_path}")
    log.info(f"Vectors: {vectors_path}")

    results = []
    for i, fp in enumerate(files, 1):
        log.info(f"\n{'='*60}")
        log.info(f"[{i}/{len(files)}] {fp.name}")
        log.info(f"{'='*60}")
        result = ingest_file(fp, conn, vectors_path, args.embed_model, args.ollama_url)
        results.append(result)

    conn.close()

    # Summary
    print("\n" + "="*60)
    print("KẾT QUẢ INGESTION")
    print("="*60)
    ok  = [r for r in results if r["status"] == "indexed"]
    dup = [r for r in results if r["status"] == "duplicate"]
    err = [r for r in results if r["status"] == "error"]

    for r in ok:
        print(f"  ✓  {r['file']:<45}  {r['chunks']} chunks")
    for r in dup:
        print(f"  ⟳  {r['file']:<45}  (đã có sẵn)")
    for r in err:
        print(f"  ✗  {r['file']:<45}  {r.get('error','')[:60]}")

    print(f"\nTổng: {len(ok)} indexed  |  {len(dup)} duplicate  |  {len(err)} error")
    print(f"\nOutput files:")
    print(f"  SQLite : {db_path}")
    print(f"  Vectors: {vectors_path}")

    # Save summary
    summary_path = out_dir / "ingestion_summary.json"
    summary_path.write_text(json.dumps({
        "run_at": datetime.now(timezone.utc).isoformat(),
        "model": args.embed_model,
        "results": results,
    }, ensure_ascii=False, indent=2))
    print(f"  Summary: {summary_path}")


if __name__ == "__main__":
    main()
